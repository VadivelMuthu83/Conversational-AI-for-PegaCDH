"""
CDH Knowledge Article Registry & Parser
========================================
Handles Pega CDH knowledge articles — PDF guides, how-to docs,
Pega Academy exports, KB articles, configuration docs, and
deployment runbooks stored in KNOWLEDGE_PATH.

Article types detected:
  1. Pega CDH Configuration Guide       (PDF/DOCX)
  2. NBA Strategy Documentation          (PDF/DOCX/MD)
  3. ADM Model Configuration Guide       (PDF/DOCX)
  4. Pega KB Article                     (PDF/HTML/TXT)
  5. Data Dictionary / Schema Doc        (XLSX/CSV/MD)
  6. Deployment / Release Notes          (MD/TXT/DOCX)
  7. Engagement Policy Rules             (XLSX/CSV)
  8. Custom / Unclassified               (any)

Each article is parsed into chunks with rich metadata so the RAG
retriever can distinguish "knowledge from the IH data" vs
"knowledge from the configuration guide".
"""
import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)


# ─── Article type definitions ────────────────────────────────────────────────

@dataclass
class KBArticleType:
    type_id: str
    display_name: str
    description: str
    filename_patterns: List[str]       # glob patterns on filename
    content_keywords: List[str]        # keywords found in content (first 500 chars)
    chunk_strategy: str = "recursive"
    tags: List[str] = field(default_factory=list)
    summary_prefix: str = ""           # prepended to chunk text for LLM context


KB_ARTICLE_TYPES: Dict[str, KBArticleType] = {

    "cdh_config_guide": KBArticleType(
        type_id="cdh_config_guide",
        display_name="CDH Configuration Guide",
        description="Pega Customer Decision Hub setup, configuration, and deployment documentation",
        filename_patterns=["*CDH*Config*", "*CustomerDecisionHub*", "*cdh_setup*",
                           "*pega*config*", "*CDH*Guide*", "*cdh_deploy*"],
        content_keywords=["customer decision hub", "cdh", "configuration", "pega platform",
                          "deployment", "environment", "tenant"],
        chunk_strategy="recursive",
        tags=["cdh", "configuration", "setup", "pega"],
        summary_prefix="[CDH Configuration Guide] ",
    ),

    "nba_strategy_doc": KBArticleType(
        type_id="nba_strategy_doc",
        display_name="NBA Strategy Documentation",
        description="Next Best Action strategy design, engagement policies, and arbitration rules",
        filename_patterns=["*NBA*", "*NextBestAction*", "*Strategy*", "*EngagementPolicy*",
                           "*nba_strategy*", "*arbitration*", "*engagement_policy*"],
        content_keywords=["next best action", "nba", "strategy", "engagement policy",
                          "arbitration", "priority", "propensity", "action", "issue", "group"],
        chunk_strategy="recursive",
        tags=["nba", "strategy", "engagement", "arbitration", "policy"],
        summary_prefix="[NBA Strategy Doc] ",
    ),

    "adm_guide": KBArticleType(
        type_id="adm_guide",
        display_name="ADM Configuration Guide",
        description="Adaptive Decision Manager model setup, predictor configuration, and tuning guides",
        filename_patterns=["*ADM*Guide*", "*AdaptiveModel*", "*adm_config*",
                           "*predictor*", "*model_config*", "*ADM*Setup*"],
        content_keywords=["adaptive decision manager", "adm", "predictor", "model",
                          "champion", "challenger", "auc", "gradient boost", "naive bayes"],
        chunk_strategy="recursive",
        tags=["adm", "model", "predictor", "configuration"],
        summary_prefix="[ADM Guide] ",
    ),

    "pega_kb_article": KBArticleType(
        type_id="pega_kb_article",
        display_name="Pega KB Article",
        description="Pega Knowledge Base articles, how-to guides, and troubleshooting documents",
        filename_patterns=["*KB-*", "*kb_*", "*KnowledgeBase*", "*PegaKB*",
                           "*HOW-TO*", "*how_to*", "*troubleshoot*", "*PRKB*"],
        content_keywords=["pega", "knowledge base", "resolution", "workaround",
                          "error", "issue", "procedure", "step", "navigate"],
        chunk_strategy="recursive",
        tags=["kb", "pega", "knowledge-base", "how-to"],
        summary_prefix="[Pega KB Article] ",
    ),

    "data_dictionary": KBArticleType(
        type_id="data_dictionary",
        display_name="Data Dictionary / Schema Doc",
        description="Column definitions, field mappings, data dictionary for CDH data exports",
        filename_patterns=["*DataDictionary*", "*data_dictionary*", "*Schema*",
                           "*FieldMapping*", "*field_map*", "*column_def*"],
        content_keywords=["field", "column", "data type", "description", "mapping",
                          "property", "class", "pxinteractionid", "pyworkid"],
        chunk_strategy="fixed",
        tags=["schema", "data-dictionary", "mapping", "fields"],
        summary_prefix="[Data Dictionary] ",
    ),

    "release_notes": KBArticleType(
        type_id="release_notes",
        display_name="Release Notes / Deployment Docs",
        description="Release notes, change logs, and deployment runbooks for CDH versions",
        filename_patterns=["*Release*Notes*", "*release_notes*", "*CHANGELOG*",
                           "*changelog*", "*RunBook*", "*runbook*", "*deployment*notes*"],
        content_keywords=["release", "version", "change", "fix", "enhancement",
                          "deploy", "migration", "upgrade", "breaking"],
        chunk_strategy="recursive",
        tags=["release", "changelog", "deployment", "version"],
        summary_prefix="[Release Notes] ",
    ),

    "engagement_policy": KBArticleType(
        type_id="engagement_policy",
        display_name="Engagement Policy Rules",
        description="Business rules, eligibility conditions, and suppression rules for NBA actions",
        filename_patterns=["*Engagement*Policy*", "*eligibility*", "*suppression*",
                           "*business_rules*", "*action_rules*", "*contact_policy*"],
        content_keywords=["eligibility", "suppression", "contact policy", "rule",
                          "condition", "when", "filter", "scope", "priority"],
        chunk_strategy="recursive",
        tags=["engagement-policy", "eligibility", "rules", "suppression"],
        summary_prefix="[Engagement Policy] ",
    ),
}


def detect_article_type(filename: str, content_preview: str = "") -> Optional[KBArticleType]:
    """
    Detect knowledge article type from filename pattern + content keywords.
    Returns None if not a knowledge article (likely raw data file).
    """
    import fnmatch
    fname = Path(filename).stem       # no extension
    ext   = Path(filename).suffix.lower()

    # Only classify document-type files as KB articles
    doc_extensions = {".pdf", ".docx", ".doc", ".html", ".htm", ".md", ".txt"}
    # Tabular files are data, not knowledge articles (unless name matches strongly)
    data_extensions = {".csv", ".tsv", ".parquet", ".json", ".jsonl"}

    for art_type in KB_ARTICLE_TYPES.values():
        # Pattern match on filename
        for pattern in art_type.filename_patterns:
            pat = pattern.lower().replace("*", "")
            if pat in fname.lower():
                return art_type

        # For doc files: also match on content keywords
        if ext in doc_extensions and content_preview:
            preview_lower = content_preview.lower()
            keyword_hits = sum(1 for kw in art_type.content_keywords if kw in preview_lower)
            if keyword_hits >= 2:
                return art_type

    return None


# ─── Knowledge article parser ────────────────────────────────────────────────

class KnowledgeArticleParser:
    """
    Parses Pega CDH knowledge articles (PDF, DOCX, MD, HTML, TXT)
    into richly-tagged chunks for RAG retrieval.

    Key difference from raw data parsing:
    - Extracts section headings and preserves structure
    - Tags each chunk with article_type, section, source_url
    - Generates a metadata header prepended to each chunk
      so the LLM knows it's reading a KB article vs IH data
    """

    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 150):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def parse(
        self,
        content: bytes,
        filename: str,
        source_url: Optional[str] = None,
    ) -> List[Dict]:
        """
        Returns list of chunk dicts:
        {
          "text":         str,    # chunk text with metadata header
          "metadata": {
            "filename":     str,
            "article_type": str,
            "section":      str,
            "chunk_index":  int,
            "source_url":   str | None,
            "tags":         list[str],
            "is_kb_article": True,
          }
        }
        """
        ext = Path(filename).suffix.lower()

        # Extract raw text
        raw_text = self._extract_text(content, filename, ext)

        # Detect article type
        preview = raw_text[:500]
        art_type = detect_article_type(filename, preview)

        if art_type is None:
            art_type = KBArticleType(
                type_id="general_doc",
                display_name="General Document",
                description="Unclassified CDH-related document",
                filename_patterns=[],
                content_keywords=[],
                tags=["general"],
                summary_prefix="[Document] ",
            )

        logger.info(f"Knowledge article: {filename} → {art_type.display_name}")

        # Extract sections
        sections = self._extract_sections(raw_text)

        # Build tagged chunks
        chunks = []
        chunk_index = 0

        for section_title, section_text in sections:
            # Prepend metadata header to each chunk
            for text_chunk in self._chunk_text(section_text):
                header = (
                    f"{art_type.summary_prefix}"
                    f"File: {filename}"
                    + (f" | Section: {section_title}" if section_title else "")
                    + "\n"
                )
                tagged_text = header + text_chunk

                chunks.append({
                    "text": tagged_text,
                    "metadata": {
                        "filename":     filename,
                        "article_type": art_type.type_id,
                        "article_name": art_type.display_name,
                        "section":      section_title or "General",
                        "chunk_index":  chunk_index,
                        "source_url":   source_url,
                        "tags":         art_type.tags + ["kb_article"],
                        "is_kb_article": True,
                    },
                })
                chunk_index += 1

        logger.info(f"  → {len(chunks)} chunks from {len(sections)} sections")
        return chunks

    def _extract_text(self, content: bytes, filename: str, ext: str) -> str:
        if ext == ".pdf":
            return self._extract_pdf(content, filename)
        elif ext in (".docx", ".doc"):
            return self._extract_docx(content, filename)
        elif ext in (".html", ".htm"):
            return self._extract_html(content)
        else:
            return content.decode("utf-8", errors="replace")

    def _extract_pdf(self, content: bytes, filename: str) -> str:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i+1}]\n{text}")
            return "\n\n".join(pages)
        except ImportError:
            try:
                import pdfminer.high_level as pm
                return pm.extract_text(io.BytesIO(content))
            except ImportError:
                logger.warning("No PDF library. pip install pypdf")
                return f"[PDF unavailable: {filename}]"
        except Exception as e:
            return f"[PDF error: {e}]"

    def _extract_docx(self, content: bytes, filename: str) -> str:
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    style = para.style.name if para.style else ""
                    if "Heading 1" in style:
                        parts.append(f"\n# {para.text}")
                    elif "Heading 2" in style:
                        parts.append(f"\n## {para.text}")
                    elif "Heading 3" in style:
                        parts.append(f"\n### {para.text}")
                    else:
                        parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            return "\n".join(parts)
        except ImportError:
            logger.warning("python-docx not installed. pip install python-docx")
            return f"[DOCX unavailable: {filename}]"
        except Exception as e:
            return f"[DOCX error: {e}]"

    def _extract_html(self, content: bytes) -> str:
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, "html.parser")
            for tag in soup(["script", "style", "nav", "footer"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            text = re.sub(r"<[^>]+>", " ", content.decode("utf-8", errors="replace"))
            return re.sub(r"\s+", " ", text).strip()

    def _extract_sections(self, text: str) -> List[tuple]:
        """
        Split text on Markdown/document headings.
        Returns list of (section_title, section_text) tuples.
        """
        # Match # Heading, ## Heading, [Page N], or ALL CAPS lines
        heading_pattern = re.compile(
            r"^(#{1,4}\s+.+|={3,}|[-]{3,}|\[Page \d+\]|[A-Z][A-Z\s]{8,}:?)$",
            re.MULTILINE,
        )

        parts = heading_pattern.split(text)

        if len(parts) <= 1:
            return [("", text)]

        sections = []
        current_title = ""
        current_text = ""

        for part in parts:
            if heading_pattern.match(part.strip()):
                if current_text.strip():
                    sections.append((current_title, current_text.strip()))
                current_title = part.strip().lstrip("#").strip()
                current_text = ""
            else:
                current_text += part

        if current_text.strip():
            sections.append((current_title, current_text.strip()))

        return sections if sections else [("", text)]

    def _chunk_text(self, text: str) -> List[str]:
        if not text:
            return []
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunks.append(text[start:end])
            start += self.chunk_size - self.chunk_overlap
        return chunks
