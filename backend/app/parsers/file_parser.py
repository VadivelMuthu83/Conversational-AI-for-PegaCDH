"""
File parser: auto-detects and parses csv, json, jsonl, xlsx, parquet,
zip, txt, md, log, tsv, xml, yaml — AND pdf, docx, html.
Returns List[ParsedFile] with text chunks + optional structured data.
"""
import io
import json
import logging
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


def _try_import(module: str):
    try:
        import importlib
        return importlib.import_module(module)
    except ImportError:
        return None


class ParsedFile:
    def __init__(self, name: str, file_type: str, text_chunks: List[str],
                 structured_data: Optional[Any] = None, row_count: Optional[int] = None,
                 columns: Optional[List[str]] = None, summary: Optional[str] = None):
        self.name = name
        self.file_type = file_type
        self.text_chunks = text_chunks
        self.structured_data = structured_data
        self.row_count = row_count
        self.columns = columns
        self.summary = summary

    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "file_type": self.file_type,
            "chunk_count": len(self.text_chunks),
            "row_count": self.row_count,
            "columns": self.columns,
            "summary": self.summary,
        }


class FileParser:
    SUPPORTED_EXTENSIONS = {
        # Tabular
        ".csv", ".tsv", ".xlsx", ".xls", ".parquet",
        # Document / text
        ".txt", ".md", ".log", ".xml", ".yaml", ".yml",
        # Structured
        ".json", ".jsonl",
        # Archive
        ".zip",
        # Rich documents (new)
        ".pdf", ".docx", ".doc", ".html", ".htm",
    }

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 150):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def can_parse(self, filename: str) -> bool:
        return Path(filename).suffix.lower() in self.SUPPORTED_EXTENSIONS

    def parse_bytes(self, content: bytes, filename: str) -> List[ParsedFile]:
        ext = Path(filename).suffix.lower()
        try:
            if ext == ".zip":
                return self._parse_zip(content, filename)
            elif ext in (".csv", ".tsv"):
                return [self._parse_csv(content, filename)]
            elif ext in (".xlsx", ".xls"):
                return [self._parse_excel(content, filename)]
            elif ext == ".parquet":
                return [self._parse_parquet(content, filename)]
            elif ext == ".json":
                return [self._parse_json(content, filename)]
            elif ext == ".jsonl":
                return [self._parse_jsonl(content, filename)]
            elif ext == ".pdf":
                return [self._parse_pdf(content, filename)]
            elif ext in (".docx", ".doc"):
                return [self._parse_docx(content, filename)]
            elif ext in (".html", ".htm"):
                return [self._parse_html(content, filename)]
            else:
                return [self._parse_text(content, filename)]
        except Exception as e:
            logger.error(f"Failed to parse {filename}: {e}")
            return [ParsedFile(name=filename, file_type="error",
                               text_chunks=[f"Error parsing {filename}: {e}"],
                               summary=f"Parse error: {e}")]

    # ── PDF ──────────────────────────────────────────────────────────────────

    def _parse_pdf(self, content: bytes, filename: str) -> ParsedFile:
        text = ""
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(f"[Page {i+1}]\n{page_text}")
            text = "\n\n".join(pages)
            logger.info(f"PDF parsed: {filename} — {len(reader.pages)} pages, {len(text)} chars")
        except ImportError:
            try:
                import pdfminer.high_level as pdfminer
                text = pdfminer.extract_text(io.BytesIO(content))
                logger.info(f"PDF parsed via pdfminer: {filename} — {len(text)} chars")
            except ImportError:
                logger.warning(f"No PDF library found for {filename}. "
                               f"Install: pip install pypdf")
                text = f"[PDF parsing unavailable — install pypdf]\nFilename: {filename}"
        except Exception as e:
            logger.error(f"PDF parse error {filename}: {e}")
            text = f"[PDF parse error: {e}]"

        return ParsedFile(
            name=filename,
            file_type="pdf",
            text_chunks=self._chunk_text(text),
            summary=f"PDF | {len(text)} chars",
        )

    # ── DOCX ─────────────────────────────────────────────────────────────────

    def _parse_docx(self, content: bytes, filename: str) -> ParsedFile:
        text = ""
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            sections = []
            for para in doc.paragraphs:
                if para.text.strip():
                    style = para.style.name if para.style else ""
                    if "Heading" in style:
                        sections.append(f"\n## {para.text}")
                    else:
                        sections.append(para.text)
            # Extract tables
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append(" | ".join(cell.text.strip() for cell in row.cells))
                sections.append("\n".join(rows))
            text = "\n".join(sections)
            logger.info(f"DOCX parsed: {filename} — {len(text)} chars")
        except ImportError:
            logger.warning(f"python-docx not installed. Install: pip install python-docx")
            text = f"[DOCX parsing unavailable — install python-docx]\nFilename: {filename}"
        except Exception as e:
            logger.error(f"DOCX parse error {filename}: {e}")
            text = f"[DOCX parse error: {e}]"

        return ParsedFile(
            name=filename,
            file_type="docx",
            text_chunks=self._chunk_text(text),
            summary=f"Word Document | {len(text)} chars",
        )

    # ── HTML ─────────────────────────────────────────────────────────────────

    def _parse_html(self, content: bytes, filename: str) -> ParsedFile:
        text = ""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, "html.parser")
            # Remove script/style noise
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
        except ImportError:
            # Fallback: strip HTML tags with regex
            import re
            text = re.sub(r"<[^>]+>", " ", content.decode("utf-8", errors="replace"))
            text = re.sub(r"\s+", " ", text).strip()
        except Exception as e:
            text = content.decode("utf-8", errors="replace")

        return ParsedFile(
            name=filename,
            file_type="html",
            text_chunks=self._chunk_text(text),
            summary=f"HTML | {len(text)} chars",
        )

    # ── ZIP ──────────────────────────────────────────────────────────────────

    def _parse_zip(self, content: bytes, filename: str) -> List[ParsedFile]:
        results = []
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            for member in zf.infolist():
                if member.is_dir():
                    continue
                try:
                    member_bytes = zf.read(member.filename)
                    parsed = self.parse_bytes(member_bytes, member.filename)
                    results.extend(parsed)
                except Exception as e:
                    logger.warning(f"Skipping {member.filename} in zip: {e}")
        return results

    # ── Tabular ──────────────────────────────────────────────────────────────

    def _parse_csv(self, content: bytes, filename: str) -> ParsedFile:
        pd = _try_import("pandas")
        if pd is None:
            return self._parse_text(content, filename)
        sep = "\t" if filename.endswith(".tsv") else ","
        df = pd.read_csv(io.BytesIO(content), sep=sep, low_memory=False)
        return self._df_to_parsed(df, filename, "csv")

    def _parse_excel(self, content: bytes, filename: str) -> ParsedFile:
        pd = _try_import("pandas")
        if pd is None:
            return ParsedFile(name=filename, file_type="excel",
                              text_chunks=["[Excel parsing requires pandas]"])
        df = pd.read_excel(io.BytesIO(content))
        return self._df_to_parsed(df, filename, "excel")

    def _parse_parquet(self, content: bytes, filename: str) -> ParsedFile:
        pd = _try_import("pandas")
        if pd is None:
            return ParsedFile(name=filename, file_type="parquet",
                              text_chunks=["[Parquet parsing requires pandas]"])
        df = pd.read_parquet(io.BytesIO(content))
        return self._df_to_parsed(df, filename, "parquet")

    def _parse_json(self, content: bytes, filename: str) -> ParsedFile:
        text = content.decode("utf-8", errors="replace")
        try:
            data = json.loads(text)
        except Exception:
            return self._parse_text(content, filename)
        pd = _try_import("pandas")
        if pd and isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
            df = pd.DataFrame(data)
            return self._df_to_parsed(df, filename, "json")
        pretty = json.dumps(data, indent=2)
        return ParsedFile(name=filename, file_type="json",
                          text_chunks=self._chunk_text(pretty),
                          summary=f"JSON | {len(str(data))} chars")

    def _parse_jsonl(self, content: bytes, filename: str) -> ParsedFile:
        pd = _try_import("pandas")
        lines = content.decode("utf-8", errors="replace").strip().split("\n")
        rows = []
        for line in lines:
            try:
                rows.append(json.loads(line))
            except Exception:
                pass
        if pd and rows and isinstance(rows[0], dict):
            df = pd.DataFrame(rows)
            return self._df_to_parsed(df, filename, "jsonl")
        return ParsedFile(name=filename, file_type="jsonl",
                          text_chunks=self._chunk_text("\n".join(lines)))

    def _parse_text(self, content: bytes, filename: str) -> ParsedFile:
        text = content.decode("utf-8", errors="replace")
        ext = Path(filename).suffix.lower().lstrip(".") or "text"
        return ParsedFile(name=filename, file_type=ext,
                          text_chunks=self._chunk_text(text),
                          summary=f"Text | {len(text)} chars")

    # ── DataFrame helper ─────────────────────────────────────────────────────

    def _df_to_parsed(self, df, filename: str, file_type: str) -> ParsedFile:
        row_count = len(df)
        columns = list(df.columns)
        summary_parts = [
            f"File: {filename}", f"Type: {file_type}",
            f"Rows: {row_count}", f"Columns: {', '.join(str(c) for c in columns)}",
        ]
        try:
            summary_parts.append(f"\nStatistics:\n{df.describe(include='all').to_string()}")
        except Exception:
            pass
        sample_text = df.head(50).to_string(index=False)
        full_text = "\n".join(summary_parts) + f"\n\nSample rows:\n{sample_text}"
        return ParsedFile(
            name=filename, file_type=file_type,
            text_chunks=self._chunk_text(full_text),
            structured_data=df.head(100).to_dict(orient="records"),
            row_count=row_count, columns=columns,
            summary=f"{file_type.upper()} | {row_count} rows × {len(columns)} cols",
        )

    # ── Chunking ─────────────────────────────────────────────────────────────

    def _chunk_text(self, text: str) -> List[str]:
        if not text:
            return []
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunks.append(text[start:end])
            start += self.chunk_size - self.chunk_overlap
        return chunks
